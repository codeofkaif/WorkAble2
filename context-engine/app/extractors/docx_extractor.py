import io
from typing import Any
import docx

class DocxExtractor:
    @staticmethod
    def extract_text(file_path_or_bytes: Any) -> str:
        try:
            if isinstance(file_path_or_bytes, bytes):
                doc = docx.Document(io.BytesIO(file_path_or_bytes))
            else:
                doc = docx.Document(file_path_or_bytes)
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        except Exception:
            return ""

docx_extractor = DocxExtractor()
